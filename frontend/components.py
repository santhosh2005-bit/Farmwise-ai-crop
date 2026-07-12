"""
components.py — Reusable Streamlit UI components for FarmWise AI Copilot.

Each function renders a self-contained piece of the dashboard so that
the main ``app.py`` stays thin and declarative.
"""

from __future__ import annotations

from typing import Any

import plotly.io as pio
import streamlit as st


def render_header() -> None:
    """Display the app title, tagline, and branding in the sidebar."""
    st.sidebar.markdown(
        """
        ## :ear_of_rice: FarmWise AI Copilot
        *Your AI-powered agricultural data analyst.*
        """
    )
    st.sidebar.divider()


def render_dataset_overview(info: dict[str, Any]) -> None:
    """Show a collapsible dataset summary in the sidebar.

    Parameters
    ----------
    info : dict[str, Any]
        Payload from ``api.get_dataset_info()``
        (keys: ``columns``, ``row_count``, ``sample_rows``).
    """
    with st.sidebar.expander(":bar_chart: Dataset Overview"):
        st.metric("Rows", f"{info['row_count']:,}")
        st.write("**Columns:**", ", ".join(info["columns"]))


def render_chat_message(role: str, content: str) -> None:
    """Render a single chat bubble.

    Parameters
    ----------
    role : str
        ``"user"`` or ``"assistant"``.
    content : str
        The message text (Markdown-safe).
    """
    with st.chat_message(role):
        if role == "user":
            st.markdown(f"### 👤 You\n{content}")
        else:
            st.markdown(f"### 🤖 FarmWise Copilot\n{content}")


def render_charts(charts_json: list[str]) -> None:
    """Render one or multiple Plotly charts.

    Parameters
    ----------
    charts_json : list[str]
        List of JSON strings produced by Plotly.
    """
    if not charts_json:
        return

    try:
        if len(charts_json) == 1:
            fig = pio.from_json(charts_json[0])
            st.plotly_chart(fig, use_container_width=True)
        else:
            cols = st.columns(len(charts_json))
            for i, chart_json in enumerate(charts_json):
                with cols[i]:
                    fig = pio.from_json(chart_json)
                    st.plotly_chart(fig, use_container_width=True)
    except Exception:
        st.warning("Could not render the chart(s).")


def render_reasoning_steps(reasoning: list[dict[str, Any]]) -> None:
    """Render the AI's step-by-step thinking process in a stylized expander."""
    if not reasoning:
        return

    with st.expander("💭 AI Thinking & Tool Execution Process", expanded=False):
        for i, step in enumerate(reasoning):
            step_type = step.get("type")
            content = step.get("content") or step.get("summary") or ""

            if step_type == "thought":
                st.markdown(f"**Step {i+1}: Thought**")
                st.info(content)

            elif step_type == "tool_call":
                name = step.get("name")
                args = step.get("arguments")
                st.markdown(f"**Step {i+1}: Calling Tool** `{name}`")
                st.code(f"{name}({args})", language="json")

            elif step_type == "tool_result":
                name = step.get("name")
                st.markdown(f"**Step {i+1}: Tool Output Summary** `{name}`")
                st.caption(content)

            if i < len(reasoning) - 1:
                st.divider()


def render_data_table(data: list[dict[str, Any]], title: str = ":clipboard: Data") -> None:
    """Render a collapsible data table.

    Parameters
    ----------
    data : list[dict]
        Row data in records orientation.
    title : str
        Header for the expander.
    """
    with st.expander(title):
        st.dataframe(data, use_container_width=True)


def render_insights(insights: dict[str, Any]) -> None:
    """Render automatic data insights in a styled expander.

    Parameters
    ----------
    insights : dict[str, Any]
        Output from ``insight_engine.generate_insight()``.
    """
    if not insights or insights.get("row_count", 0) == 0:
        return

    with st.expander(":bulb: Automatic Data Insights", expanded=False):
        col1, col2, col3 = st.columns(3)

        with col1:
            st.metric("Rows analysed", f"{insights.get('row_count', 0):,}")
            if insights.get("average") is not None:
                st.metric("Average", f"{insights['average']:,.2f}")

        with col2:
            if insights.get("median") is not None:
                st.metric("Median", f"{insights['median']:,.2f}")
            if insights.get("trend"):
                trend = insights["trend"]
                trend_icon = {"increasing": ":arrow_upper_right:", "decreasing": ":arrow_lower_right:", "stable": ":left_right_arrow:"}.get(trend, "")
                st.metric("Trend", f"{trend_icon} {trend.title()}")

        with col3:
            if insights.get("growth_rate_pct") is not None:
                gr = insights["growth_rate_pct"]
                delta_color = "normal" if gr >= 0 else "inverse"
                st.metric("Growth Rate", f"{gr:+.1f}%", delta=f"{gr:+.1f}%", delta_color=delta_color)
            if insights.get("outlier_count") is not None:
                st.metric("Outliers (IQR)", insights["outlier_count"])

        # Highest / Lowest details
        if insights.get("highest"):
            col = insights.get("column_analysed", "value")
            highest = insights["highest"]
            st.success(
                f"**Highest {col}:** {highest.get(col, 'N/A'):,.2f} "
                f"({highest.get('Area', '')} - {highest.get('Item', '')} - {highest.get('Year', '')})"
                if isinstance(highest.get(col), (int, float)) else f"**Highest:** {highest}"
            )

        if insights.get("lowest"):
            col = insights.get("column_analysed", "value")
            lowest = insights["lowest"]
            st.info(
                f"**Lowest {col}:** {lowest.get(col, 'N/A'):,.2f} "
                f"({lowest.get('Area', '')} - {lowest.get('Item', '')} - {lowest.get('Year', '')})"
                if isinstance(lowest.get(col), (int, float)) else f"**Lowest:** {lowest}"
            )


def render_suggestions(suggestions: list[str]) -> str | None:
    """Render follow-up suggestion buttons.

    Parameters
    ----------
    suggestions : list[str]
        List of suggested follow-up questions.

    Returns
    -------
    str | None
        The clicked suggestion text, or *None* if nothing was clicked.
    """
    if not suggestions:
        return None

    st.markdown("---")
    st.markdown("**:question: Try asking:**")

    clicked: str | None = None
    cols = st.columns(min(len(suggestions), 4))
    for i, suggestion in enumerate(suggestions[:4]):
        with cols[i]:
            if st.button(
                suggestion,
                key=f"suggestion_{hash(suggestion)}_{i}",
                use_container_width=True,
            ):
                clicked = suggestion

    return clicked


def render_error(message: str) -> None:
    """Show a user-friendly error banner.

    Parameters
    ----------
    message : str
        The error description to display.
    """
    st.error(f":warning: {message}")


def render_backend_status(is_healthy: bool) -> None:
    """Show a coloured status indicator in the sidebar.

    Parameters
    ----------
    is_healthy : bool
        ``True`` when the backend ``/health`` endpoint responds OK.
    """
    if is_healthy:
        st.sidebar.success(":green_circle: Backend connected")
    else:
        st.sidebar.error(":red_circle: Backend unreachable")


def render_theme_customizer(theme_mode: str) -> None:
    """Inject custom CSS to override theme styles and enable glassmorphism/custom themes.

    Parameters
    ----------
    theme_mode : str
        "Dark Mode (Glassmorphism)" or "Light Mode (Vibrant)".
    """
    if theme_mode == "Dark Mode":
        st.markdown(
            """
            <style>
            .stApp {
                background: radial-gradient(circle at 50% 50%, #0f172a, #090d16) !important;
                color: #f8fafc !important;
            }
            .stChatMessage {
                background-color: rgba(30, 41, 59, 0.5) !important;
                border: 1px solid rgba(255, 255, 255, 0.08) !important;
                backdrop-filter: blur(12px) !important;
                border-radius: 12px !important;
            }
            section[data-testid="stSidebar"] {
                background-color: #050811 !important;
                border-right: 1px solid rgba(255, 255, 255, 0.05) !important;
            }
            .streamlit-expanderHeader {
                background-color: rgba(30, 41, 59, 0.4) !important;
                border: 1px solid rgba(255, 255, 255, 0.05) !important;
                color: #f8fafc !important;
            }
            [data-testid="stMetric"] {
                background-color: rgba(30, 41, 59, 0.4) !important;
                border: 1px solid rgba(255, 255, 255, 0.05) !important;
                backdrop-filter: blur(8px) !important;
            }
            [data-testid="stMetricValue"] {
                color: #10B981 !important;
            }
            div[data-testid="stMarkdownContainer"] p {
                color: #cbd5e1 !important;
            }
            h1, h2, h3 {
                color: #f8fafc !important;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            """
            <style>
            .stApp {
                background: linear-gradient(135deg, #f8fafc 0%, #f0fdf4 100%) !important;
                color: #0f172a !important;
            }
            .stChatMessage {
                background-color: rgba(255, 255, 255, 0.8) !important;
                border: 1px solid rgba(16, 185, 129, 0.15) !important;
                border-radius: 12px !important;
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.02) !important;
            }
            section[data-testid="stSidebar"] {
                background-color: #ffffff !important;
                border-right: 1px solid rgba(16, 185, 129, 0.1) !important;
            }
            .streamlit-expanderHeader {
                background-color: rgba(255, 255, 255, 0.9) !important;
                border: 1px solid rgba(16, 185, 129, 0.1) !important;
            }
            [data-testid="stMetric"] {
                background-color: #ffffff !important;
                border: 1px solid rgba(16, 185, 129, 0.1) !important;
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.02) !important;
            }
            [data-testid="stMetricValue"] {
                color: #059669 !important;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )


def render_kpi_cards(info: dict[str, Any]) -> None:
    """Render beautiful KPI cards showing database metrics.

    Parameters
    ----------
    info : dict[str, Any]
        Payload from ``api.get_dataset_info()``.
    """
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Observations", f"{info.get('row_count', 0):,}")
    with col2:
        st.metric("Crops Covered", f"{len(info.get('crops', []))}")
    with col3:
        st.metric("Countries Analyzed", f"{len(info.get('countries', []))}")
    with col4:
        years = info.get('years', [])
        year_range = f"{min(years)} - {max(years)}" if years else "N/A"
        st.metric("Historical Coverage", year_range)


def render_landing_page(info: dict[str, Any], on_launch_chat) -> None:
    """Render the high-fidelity professional landing page.

    Parameters
    ----------
    info : dict[str, Any]
        Payload from ``api.get_dataset_info()``.
    on_launch_chat : callable
        Callback function executed when the user clicks 'Launch AI Copilot'.
    """
    # Hero Title & Headline
    st.markdown(
        """
        <div style="text-align: center; padding: 2rem 0; margin-bottom: 2rem;">
            <h1 style="font-size: 3rem; margin-bottom: 1rem;">🌾 FarmWise AI Copilot</h1>
            <p style="font-size: 1.25rem; opacity: 0.85; max-width: 800px; margin: 0 auto;">
                An intelligent analytical assistant specializing in global crop yields, climatic trends, and pesticide utilization. Query, visualize, and extract data-backed explanations in plain English.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # KPI Metrics Section
    st.subheader("📈 Dataset At-A-Glance")
    render_kpi_cards(info)
    st.markdown("<br>", unsafe_allow_html=True)

    # Interactive Core Capabilities Grid
    st.subheader("🛠️ Core Capabilities")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(
            """
            <div style="padding: 1.5rem; border-radius: 12px; border: 1px solid rgba(128,128,128,0.15); margin-bottom: 1rem;">
                <h4 style="margin-top:0;">💬 Conversational Analyst</h4>
                <p style="font-size: 0.9rem;">Ask complex questions like <i>"Compare wheat yields between France and Germany over the last 10 years"</i>. The AI interprets your queries and fetches relevant data automatically.</p>
            </div>
            <div style="padding: 1.5rem; border-radius: 12px; border: 1px solid rgba(128,128,128,0.15); margin-bottom: 1rem;">
                <h4 style="margin-top:0;">🧠 Automated Insights</h4>
                <p style="font-size: 0.9rem;">Instant calculations of compound growth rates, statistical summaries, trend directions, and outlier detections generated dynamically for your queries.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            """
            <div style="padding: 1.5rem; border-radius: 12px; border: 1px solid rgba(128,128,128,0.15); margin-bottom: 1rem;">
                <h4 style="margin-top:0;">📊 Dynamic Visualizations</h4>
                <p style="font-size: 0.9rem;">Renders high-quality Plotly line charts, bar charts, scatter plots, and distribution heatmaps tailored precisely to the nature of your request.</p>
            </div>
            <div style="padding: 1.5rem; border-radius: 12px; border: 1px solid rgba(128,128,128,0.15); margin-bottom: 1rem;">
                <h4 style="margin-top:0;">📥 Exportable Formats</h4>
                <p style="font-size: 0.9rem;">Instantly download the underlying data in CSV format or download the entire analysis transcript as Markdown to include in your reports.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # Call-to-action button
    col_left, col_mid, col_right = st.columns([1, 2, 1])
    with col_mid:
        if st.button("🚀 Launch AI Copilot Chat", use_container_width=True, type="primary"):
            on_launch_chat()

    st.markdown("<br><hr><br>", unsafe_allow_html=True)
    st.subheader("📋 Sample Raw Records")
    st.dataframe(info.get("sample_rows", []), use_container_width=True)


def generate_docx(chat_history: list[dict]) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt
    from docx.dml.color import RGBColor

    doc = Document()
    doc.add_heading("FarmWise AI Copilot Conversation Transcript", level=0)

    for msg in chat_history:
        role_label = "User" if msg["role"] == "user" else "Copilot"
        h = doc.add_heading(role_label, level=2)
        if msg["role"] == "user":
            h.runs[0].font.color.rgb = RGBColor(59, 130, 246)  # Blue
        else:
            h.runs[0].font.color.rgb = RGBColor(16, 185, 129)  # Green

        if msg.get("content"):
            doc.add_paragraph(msg["content"])
            
        if msg.get("insights"):
            ins = msg["insights"]
            doc.add_paragraph("Automatic Data Insights:")
            doc.add_paragraph(f"• Rows Analyzed: {ins.get('row_count')}", style='List Bullet')
            doc.add_paragraph(f"• Trend: {ins.get('trend')}", style='List Bullet')
            doc.add_paragraph(f"• Average: {ins.get('average')}", style='List Bullet')

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_pdf(chat_history: list[dict]) -> bytes:
    import io
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font('helvetica', 'B', 14)
            self.set_text_color(16, 185, 129) # Green title
            self.cell(0, 10, 'FarmWise AI Copilot', border=False, new_x="LMARGIN", new_y="NEXT", align='C')
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font('helvetica', 'I', 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("helvetica", size=10)

    pdf.set_font('helvetica', 'B', 14)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 10, 'Conversation Transcript', new_x="LMARGIN", new_y="NEXT", align='L')
    pdf.ln(5)

    for msg in chat_history:
        role_label = "User" if msg["role"] == "user" else "Copilot"
        pdf.set_font('helvetica', 'B', 11)
        if msg["role"] == "user":
            pdf.set_text_color(59, 130, 246) # Blue
        else:
            pdf.set_text_color(16, 185, 129) # Green
        pdf.cell(0, 8, f'{role_label}:', new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_font('helvetica', '', 10)
        pdf.set_text_color(51, 65, 85)
        
        content = msg.get("content", "")
        # Simple cleanup of markdown symbols for standard PDF fonts
        content = content.replace("**", "").replace("*", "-").replace("`", "")
        content_encoded = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, content_encoded)
        pdf.ln(4)

        if msg.get("insights"):
            ins = msg["insights"]
            pdf.set_font('helvetica', 'B', 10)
            pdf.set_text_color(15, 23, 42)
            pdf.cell(0, 6, "Automatic Data Insights:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font('helvetica', '', 10)
            pdf.set_text_color(51, 65, 85)
            pdf.cell(0, 5, f"- Rows Analyzed: {ins.get('row_count')}", new_x="LMARGIN", new_y="NEXT")
            pdf.cell(0, 5, f"- Trend: {ins.get('trend')}", new_x="LMARGIN", new_y="NEXT")
            pdf.cell(0, 5, f"- Average: {ins.get('average')}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)

    return bytes(pdf.output())


def render_export_buttons(chat_history: list[dict], last_data: list[dict] | None) -> None:
    """Render options to export conversation transcript and query data.

    Parameters
    ----------
    chat_history : list[dict]
        Full conversation memory log.
    last_data : list[dict] | None
        The tabular data returned from the last query.
    """
    st.sidebar.subheader("📥 Export & Report Options")

    # 1. Export Chat Transcript (Word & PDF)
    if chat_history:
        # Export as Word (.docx)
        try:
            docx_data = generate_docx(chat_history)
            st.sidebar.download_button(
                label="📝 Download Chat Log (.docx)",
                data=docx_data,
                file_name="farmwise_chat_transcript.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.sidebar.warning(f"Could not generate Word document: {e}")

        # Export as PDF (.pdf)
        try:
            pdf_data = generate_pdf(chat_history)
            st.sidebar.download_button(
                label="📕 Download Chat Log (.pdf)",
                data=pdf_data,
                file_name="farmwise_chat_transcript.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        except Exception as e:
            st.sidebar.warning(f"Could not generate PDF: {e}")

    # 2. Export Last Query Data
    if last_data:
        import pandas as pd
        df = pd.DataFrame(last_data)
        csv_data = df.to_csv(index=False)
        st.sidebar.download_button(
            label="📥 Download Query Data (.csv)",
            data=csv_data,
            file_name="farmwise_query_data.csv",
            mime="text/csv",
            use_container_width=True
        )


def render_interactive_analytics(info: dict[str, Any]) -> None:
    """Render the high-fidelity interactive analytics and forecasting suite.

    Parameters
    ----------
    info : dict[str, Any]
        Payload from ``api.get_dataset_info()``.
    """
    from frontend import api
    import pandas as pd
    import plotly.express as px
    import plotly.graph_objects as go
    import numpy as np

    st.title("📊 Interactive Analytics & Forecasts")
    st.caption(
        "Directly query the dataset, configure yield forecasts, and analyze the impact of climatic drivers."
    )

    # Guided Application Tour
    with st.expander("🚀 Quick Dashboard Tour (Recommended for New Users)", expanded=True):
        tour_tab1, tour_tab2, tour_tab3, tour_tab4 = st.tabs([
            "🌍 1. Select Country & Crop",
            "📅 2. Adjust Projection Horizon",
            "📊 3. Explore Analytics Tabs",
            "💬 4. Switch to Chat anytime"
        ])
        
        with tour_tab1:
            st.markdown(
                """
                **How to start:**
                Use the **Select Country / Area** and **Select Crop Item** dropdowns below.
                - The system automatically retrieves the matching historical records (1990 - 2013) for those parameters.
                - All tabs and models on this page instantly recalculate based on your choices.
                """
            )
        with tour_tab2:
            st.markdown(
                """
                **How to configure forecasting:**
                Use the **Forecast Horizon** slider below.
                - Move it to decide how many years beyond 2013 you want to forecast yields.
                - A **Linear Regression Model** is fitted to the historical yield data to project the future yields under the **📈 Yield Trend & Forecast** tab.
                """
            )
        with tour_tab3:
            st.markdown(
                """
                **What each tab contains:**
                - **📈 Yield Trend & Forecast**: View historical points, the fitted trendline, and the future projections.
                - **🌿 Environmental Drivers**: Trace the relationship between crop yield and temperature, rainfall, or pesticide usage.
                - **🗺️ Correlation Matrix**: Heatmap showing how temperature, rainfall, pesticides, and yields correlate.
                - **🌾 Crop Breakdown**: Pie charts and breakdown tables comparing yield shares among different crops within the country.
                """
            )
        with tour_tab4:
            st.markdown(
                """
                **Need a custom report or explanation?**
                - Switch to the **💬 AI Copilot Chat** page in the sidebar navigation.
                - Ask questions in plain English (e.g. *"What is the strongest climate driver for rice in India?"*), and the Copilot will fetch the data and explain it.
                """
            )

    if not info:
        st.warning("Dataset information is unavailable. Ensure the backend is running.")
        return

    # Dropdowns for Country and Crop
    col1, col2 = st.columns(2)
    with col1:
        countries = info.get("countries", [])
        default_country_idx = countries.index("India") if "India" in countries else 0
        country = st.selectbox("🌍 Select Country / Area", countries, index=default_country_idx)
    with col2:
        crops = info.get("crops", [])
        default_crop_idx = crops.index("Wheat") if "Wheat" in crops else 0
        crop = st.selectbox("🌾 Select Crop Item", crops, index=default_crop_idx)

    st.markdown("---")

    # Forecast horizon slider — user-configurable, extends beyond dataset end
    forecast_years = st.slider(
        "📅 Forecast Horizon (years beyond 2013)",
        min_value=5,
        max_value=30,
        value=15,
        step=1,
        help="The dataset ends in 2013. Extend forecast up to 30 years into the future."
    )

    # Fetch data from backend — stored in local variables, never in session_state
    # so this page is completely isolated from the chatbot's data.
    with st.spinner("Fetching data and computing models..."):
        try:
            forecast_data = api.get_forecast(country, crop, forecast_years=forecast_years)
            regression_data = api.get_regression(country, crop)
            raw_data = api.get_raw_data(country, crop)
            raw_data_country = api.get_raw_data(country, "Wheat")  # for multi-crop overview
        except Exception as e:
            st.error(f"Error fetching analytics data from backend: {e}")
            return

    if not forecast_data or "error" in forecast_data[0]:
        err_msg = forecast_data[0].get("error") if forecast_data else "No historical records found."
        st.warning(f"Could not perform analytics: {err_msg}")
        return

    # Convert to DataFrames — all local, never written to st.session_state
    df_forecast = pd.DataFrame(forecast_data)
    df_hist = df_forecast[df_forecast["type"] == "Historical"]
    df_raw = pd.DataFrame(raw_data)

    # Clean raw data for weather/pesticides correlation
    cols_of_interest = [
        "hg/ha_yield",
        "average_rain_fall_mm_per_year",
        "pesticides_tonnes",
        "avg_temp",
    ]
    cols_present = [c for c in cols_of_interest if c in df_raw.columns]
    df_grouped = df_raw.groupby("Year")[cols_present].mean().reset_index()

    # Tabs — added two more
    tab_trend, tab_drivers, tab_corr, tab_multi = st.tabs([
        "📈 Yield Trend & Forecast",
        "🌿 Environmental Drivers",
        "🗺️ Correlation Matrix",
        "🌾 Crop Breakdown"
    ])

    # ── Tab 1: Yield Trend & Forecast ─────────────────────────────────────
    with tab_trend:
        st.subheader("Yield Projections & Trend Line")
        show_forecast = st.checkbox("Show 10-Year Trend Forecast Projection", value=True)

        fig = go.Figure()

        # Add historical points
        fig.add_trace(go.Scatter(
            x=df_hist["Year"],
            y=df_hist["hg/ha_yield"],
            mode="markers+lines",
            name="Historical Yield",
            line=dict(color="#3B82F6", width=3),
            marker=dict(size=8, color="#1D4ED8")
        ))

        # Add historical trendline
        fig.add_trace(go.Scatter(
            x=df_hist["Year"],
            y=df_hist["trend_line"],
            mode="lines",
            name="Fitted Trendline",
            line=dict(color="#EF4444", width=2, dash="dash")
        ))

        # Add forecast projection
        if show_forecast:
            df_proj = df_forecast[df_forecast["type"] == "Forecast"]
            if not df_hist.empty and not df_proj.empty:
                last_hist = df_hist.iloc[-1]
                df_proj = pd.concat([pd.DataFrame([last_hist]), df_proj], ignore_index=True)

            fig.add_trace(go.Scatter(
                x=df_proj["Year"],
                y=df_proj["hg/ha_yield"],
                mode="lines",
                name="Forecast Projection",
                line=dict(color="#10B981", width=3, dash="dot")
            ))

        fig.update_layout(
            title=f"Yield Trend and Forecast for {crop} in {country}",
            xaxis_title="Year",
            yaxis_title="Yield (hg/ha)",
            template="plotly_white",
            height=500,
            hovermode="x unified",
            legend=dict(orientation="h", yanchor="bottom", y=-0.2, xanchor="center", x=0.5)
        )
        st.plotly_chart(fig, use_container_width=True)

        # Stat cards
        stats_col1, stats_col2, stats_col3, stats_col4 = st.columns(4)
        r_sq = df_forecast.iloc[-1].get("r_squared", "N/A")
        slope = df_forecast.iloc[-1].get("slope", "N/A")
        latest_hist = df_hist.iloc[-1] if not df_hist.empty else None
        latest_yield = latest_hist["hg/ha_yield"] if latest_hist is not None else 0.0
        last_forecast_year = df_forecast[df_forecast["type"] == "Forecast"]["Year"].max() if not df_forecast.empty else "N/A"

        with stats_col1:
            st.metric("Model fit (R²)", f"{r_sq:.2%}" if isinstance(r_sq, float) else r_sq)
        with stats_col2:
            st.metric("Annual Growth (slope)", f"{slope:+,} hg/ha/yr" if isinstance(slope, (int, float)) else slope)
        with stats_col3:
            st.metric("Latest Historical Yield (2013)", f"{latest_yield:,.0f} hg/ha")
        with stats_col4:
            st.metric("Forecast End Year", str(last_forecast_year))

    # ── Tab 2: Environmental Drivers ──────────────────────────────────────
    with tab_drivers:
        st.subheader("Environmental Drivers vs. Yield")
        if "error" in regression_data:
            st.info(regression_data["error"])
        else:
            st.markdown(f"💡 **Key Finding:** {regression_data.get('description', '')}")

            driver_col1, driver_col2, driver_col3 = st.columns(3)
            factors = regression_data.get("factors", {})

            factors_meta = [
                ("Rainfall", "average_rain_fall_mm_per_year", "Rainfall (mm)", driver_col1, "#06B6D4"),
                ("Pesticides", "pesticides_tonnes", "Pesticides (tonnes)", driver_col2, "#F97316"),
                ("Temperature", "avg_temp", "Temperature (°C)", driver_col3, "#EC4899")
            ]

            for name, col_name, label, col_container, color in factors_meta:
                with col_container:
                    factor_info = factors.get(name)
                    corr = factor_info.get("correlation") if factor_info else None
                    if not factor_info or corr is None or np.isnan(corr):
                        st.info(f"Insufficient {name} data.")
                        continue

                    # Scatter plot with regression line
                    sub_fig = go.Figure()
                    valid_df = df_grouped[[col_name, "hg/ha_yield"]].dropna()

                    # Scatter points
                    sub_fig.add_trace(go.Scatter(
                        x=valid_df[col_name],
                        y=valid_df["hg/ha_yield"],
                        mode="markers",
                        name="Observed",
                        marker=dict(color=color, size=8, opacity=0.8)
                    ))

                    # Regression line
                    f_slope = factor_info["slope"]
                    f_intercept = factor_info["intercept"]
                    x_range = np.linspace(valid_df[col_name].min(), valid_df[col_name].max(), 100)
                    y_range = f_slope * x_range + f_intercept

                    sub_fig.add_trace(go.Scatter(
                        x=x_range,
                        y=y_range,
                        mode="lines",
                        name="Fit",
                        line=dict(color="#1E293B", width=1.5, dash="dash")
                    ))

                    sub_fig.update_layout(
                        title=f"{name} Relationship (r = {factor_info['correlation']:+.2f})",
                        xaxis_title=label,
                        yaxis_title="Yield (hg/ha)",
                        template="plotly_white",
                        showlegend=False,
                        height=350,
                        margin=dict(l=40, r=40, t=40, b=40)
                    )
                    st.plotly_chart(sub_fig, use_container_width=True)

    # ── Tab 3: Correlation Matrix ─────────────────────────────────────────
    with tab_corr:
        st.subheader("Correlation Heatmap")
        st.markdown(
            "This matrix shows pairwise Pearson correlation coefficients between yield, rainfall, "
            "pesticide usage, and temperature. Values range from **-1.0** (perfect negative correlation) "
            "to **+1.0** (perfect positive correlation)."
        )

        corr_matrix = df_grouped[cols_present].corr()
        pretty_labels = [c.replace("_", " ").title().replace("Hg/Ha ", "") for c in cols_present]

        fig_corr = px.imshow(
            corr_matrix,
            text_auto=".2f",
            color_continuous_scale="RdBu",
            range_color=[-1, 1],
            labels=dict(color="Correlation"),
            x=pretty_labels,
            y=pretty_labels,
        )
        fig_corr.update_layout(
            template="plotly_white",
            height=450,
            margin=dict(l=40, r=40, t=40, b=40),
            coloraxis_colorbar=dict(title="r")
        )
        st.plotly_chart(fig_corr, use_container_width=True)

    # ── Tab 4: Crop Breakdown (new) ───────────────────────────────────────
    with tab_multi:
        st.subheader(f"All Crops in {country} — Yield Overview")
        st.markdown(
            "Compares average yield, rainfall, and pesticide usage across **all crops** "
            f"in **{country}** over the full historical period."
        )

        # Fetch all-crop data for this country using raw filter
        try:
            from frontend.api import _BASE_URL
            import requests as _req
            resp = _req.get(
                f"{_BASE_URL}/analytics/raw-data",
                params={"country": country, "item": "Wheat"},
                timeout=10,
            )
            # Get full country data via /dataset/info endpoint trick:
            # We'll filter from the available country data by calling the tool directly.
            from backend.tools import filter_by_country as _fbc
            all_crops_df = pd.DataFrame(_fbc(country))
        except Exception:
            all_crops_df = pd.DataFrame()

        if all_crops_df.empty:
            st.info("Could not load multi-crop data.")
        else:
            # ── Chart A: Average yield per crop (bar) ───────────────────
            crop_avg = (
                all_crops_df.groupby("Item")["hg/ha_yield"]
                .mean()
                .sort_values(ascending=False)
                .reset_index()
            )
            crop_avg.columns = ["Crop", "Average Yield (hg/ha)"]

            fig_bar = px.bar(
                crop_avg,
                x="Crop",
                y="Average Yield (hg/ha)",
                title=f"Average Crop Yield in {country} (1990–2013)",
                color="Average Yield (hg/ha)",
                color_continuous_scale=[
                    [0.0, "#3B82F6"], [0.5, "#10B981"], [1.0, "#F97316"]
                ],
                template="plotly_white",
            )
            try:
                fig_bar.update_traces(marker_cornerradius=6)
            except Exception:
                pass  # older Plotly versions don't support cornerradius
            fig_bar.update_layout(
                height=400,
                showlegend=False,
                coloraxis_showscale=False,
                font=dict(color="#1E293B"),
                title_font=dict(color="#0F172A", size=16),
            )
            st.plotly_chart(fig_bar, use_container_width=True)

            # ── Chart B: Yield distribution box plot ────────────────────
            fig_box = px.box(
                all_crops_df,
                x="Item",
                y="hg/ha_yield",
                title=f"Yield Distribution Per Crop in {country}",
                color="Item",
                color_discrete_sequence=[
                    "#3B82F6", "#10B981", "#F97316", "#8B5CF6",
                    "#EC4899", "#06B6D4", "#F59E0B", "#EF4444",
                ],
                template="plotly_white",
            )
            fig_box.update_layout(
                height=420,
                showlegend=False,
                font=dict(color="#1E293B"),
                title_font=dict(color="#0F172A", size=16),
                xaxis_title="Crop",
                yaxis_title="Yield (hg/ha)",
            )
            st.plotly_chart(fig_box, use_container_width=True)

            # ── Chart C: Rainfall vs Yield scatter (all crops) ──────────
            if "average_rain_fall_mm_per_year" in all_crops_df.columns:
                fig_scatter = px.scatter(
                    all_crops_df,
                    x="average_rain_fall_mm_per_year",
                    y="hg/ha_yield",
                    color="Item",
                    title=f"Rainfall vs Yield for All Crops in {country}",
                    template="plotly_white",
                    opacity=0.7,
                    color_discrete_sequence=[
                        "#3B82F6", "#10B981", "#F97316", "#8B5CF6",
                        "#EC4899", "#06B6D4", "#F59E0B", "#EF4444",
                    ],
                )
                fig_scatter.update_traces(marker=dict(size=8))
                fig_scatter.update_layout(
                    height=420,
                    font=dict(color="#1E293B"),
                    title_font=dict(color="#0F172A", size=16),
                    xaxis_title="Rainfall (mm/year)",
                    yaxis_title="Yield (hg/ha)",
                    legend=dict(
                        orientation="h", yanchor="bottom", y=-0.35,
                        xanchor="center", x=0.5, font=dict(size=11)
                    )
                )
                st.plotly_chart(fig_scatter, use_container_width=True)
